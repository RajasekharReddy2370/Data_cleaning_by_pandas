import pandas as pd
from docx import Document
import os
# from docx2pdf import convert
import pdfkit

# Read the Excel file into a DataFrame
df = pd.read_excel(r"/home/rajashekar/WORK/voters_prac/94_GTR.xlsx")

# Path to the Word document template
template_file_path = "/home/rajashekar/Downloads/FORM_NEW_12D (1).docx"

# Check if the template file exists
if not os.path.exists(template_file_path):
    print(f"Error: Template file '{template_file_path}' not found.")
else:
    try:
        # Read the Word document template
        template_doc = Document(template_file_path)

        # Iterate over each row in the DataFrame
        for index, row in df.iterrows():
            print(row['Voter_id'])
            if row['RLN_TYPE'] == 'F':
                row['RLN_TYPE'] = 'Father'
            if row['RLN_TYPE'] == 'M':
                row['RLN_TYPE'] = 'Mother'
            if row['RLN_TYPE'] == 'H':
                row['RLN_TYPE'] = 'Husband'
            # Generate form text based on DataFrame values
            form_text = f"""                                                    FORM 12D
                                                [see rule 27-C] 
                                                        PART I
                            Letter of intimation to Assistant Returning Officer 
                                                (for absentee voters)
                                                                                                                        To
                                                                                                                        The Assistant Returning Officer,
                                                                                                                        94_Guntur West Parliamentary/Assembly constituency
                                                                                                                        …………………………………... (designation and address of ARO)
                                                                                                                        Sir,
                                                                                                                            I, {row['Names']} {row['RLN_TYPE']} {row['RLN_FM_NM_EN']}, resident of {row['Address_eng']} village/mohalla Guntur Town/city/tehsil Guntur District, Andhrapradesh (State) belong to the class of absentee voter and wish to cast my vote by post at the election to the House of the People/Legislative Assembly from the 94-Guntur West Parliamentary/Assembly constituency. 
                                                                                                                        My complete present postal address is as under:-                                    House/dwelling unit/tent number {row['House No']},                              Camp/mohalla/village {row['Address_eng']}.                                                                      Ward/town/tehsil Guntur,                                                                                        District Guntur,                                                                                                State Andhrapradesh, PIN CODE 522002.                                                                  Mobile Phone No. (if available) {row['Mobile']}.                                                                                                                                                                                                                                                                                                       
                                                                                                                        My name is entered at serial number {row['S']} in part No {row['B#']} of the electoral roll for 94-Guntur West Assembly constituency.
                                                                                                                            I am working as ............... in ..............
                                                                                                                            I will be on duty in the above-mentioned office on the day of poll for the above-mentioned election. 
                                                                                                                            *On account of my official duties on the date of poll, I will not be in a position to be present in the polling station assigned to me on the day of poll.                                                                                                                          or                                                                                                                          *I am {row['Age']} years of age/am a person with disability, and am not in a position to go to the polling station to cast vote.                                                                                                                                                                                                                                                                                                                                                                                                      
                                                                                                                            It is requested that postal ballot paper may be issued to me as absentee voter for the above election.                     
                                                                                                Yours faithfully,
                                                                                {row['Names']} 
                                                                                (Full name and signature)
                                                            PART II
                                                                                                                                    (for absentee voter other than senior citizen or persons with disability)   Certificate by the nodal officer appointed by the Organization concerned.                            It is hereby certified that the particulars given by the applicant in Part I are correct, and it is further certified that the applicant will be on official duty on the day of poll, and he/she will not be in a position to be present in the polling station on the day of poll.                                                                                                                                                                                                                                                                                                                                                                                                                                                                            
                                                                        .........................................
                                                        (full signature of the attesting Officer)
                                                                   ………………………………(Name)
                                                                 …………………………….(address)
                                                                 ……………………….(rubber stamp)
                                                                                                                            • Strike off whichever is not applicable and tick the relevant statement.
                                                                                                                            Note- This Application must reach RO within 5 days following the date of notification of election."""


            # Create a new Document object for the filled form
            filled_doc = Document()

            # Add paragraphs to the new document based on the filled form text
            for para in form_text.split('\n'):
                filled_doc.add_paragraph(para)

            # Save the filled form as a Word document
            filled_file_path = f"/home/rajashekar/WORK/voters_prac/voter_id_docs/{row['B#']}-{row['Names']}.docx"
            filled_doc.save(filled_file_path)

            # pdf_file_path = f"RRR_PDF{index}.pdf"
            # pdfkit.from_file(filled_file_path, pdf_file_path)
            break
    except Exception as e:
        print(f"An error occurred: {e}")
